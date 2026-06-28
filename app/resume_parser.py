import re
from pathlib import Path
from docx import Document
from docx.document import Document as DocxDocument
from typing import Dict, List, Optional, Tuple
import logging

logger = logging.getLogger(__name__)


class ResumeParser:
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = None
        self.contact_info = {}
        self.sections = {}

        if self.docx_path.exists():
            self.doc = Document(str(self.docx_path))
            logger.info(f"Loaded resume from {self.docx_path}")
        else:
            logger.warning(f"Resume file not found: {self.docx_path}")

    def extract_contact_info(self) -> Dict:
        """Extract name, email, phone, location from the resume header."""
        if not self.doc:
            return {}

        contact = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None,
        }

        # Extract all text from first paragraph which contains contact info
        header_text = ""
        for para in self.doc.paragraphs[:10]:
            text = para.text.strip()
            if text:
                header_text += " " + text
            # Stop at first section heading
            if any(kw in text.lower() for kw in ["summary", "skills", "experience", "education"]):
                break

        # Email extraction
        email_match = re.search(r'[\w.-]+@[\w.-]+\.\w+', header_text)
        if email_match:
            contact["email"] = email_match.group()

        # Phone extraction (various formats)
        phone_match = re.search(
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            header_text
        )
        if phone_match:
            contact["phone"] = phone_match.group().strip()

        # Name extraction - first line/phrase that's not contact details
        lines = header_text.split('\n')
        for line in lines:
            line = line.strip()
            if line and '@' not in line and not re.search(r'\(\d{3}\)', line):
                # Check if it looks like a name (2-4 words, capitalized)
                words = line.split()
                if 1 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
                    contact["name"] = line
                    break

        # Location extraction (look for city, state patterns)
        location_match = re.search(
            r'([A-Za-z\s]+),\s*([A-Z]{2})|([A-Za-z\s]+),\s*([A-Za-z\s]+)',
            header_text
        )
        if location_match:
            contact["location"] = location_match.group().strip()

        self.contact_info = contact
        logger.info(f"Extracted contact info: {contact}")
        return contact

    def parse_resume(self) -> Dict:
        """Parse the entire resume into structured sections."""
        if not self.doc:
            return {}

        sections = {
            "contact": self.extract_contact_info(),
            "summary": "",
            "skills": [],
            "experience": [],
            "education": [],
            "certifications": [],
            "projects": [],
        }

        current_section = None
        current_entry = None
        skill_text = ""

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_lower = text.lower()

            # Section detection (check if this is a heading)
            # Only match short phrases that are likely section headings, not body text
            is_heading = (
                len(text) < 50 and any(
                    re.search(r'\b' + keyword + r'\b', text_lower)
                    for keyword in [
                        "summary", "profile",
                        "skills", "technical",
                        "experience", "professional experience",
                        "education", "academic",
                        "certification", "licenses", "awards",
                        "project", "portfolio"
                    ]
                )
            )

            if is_heading:
                if "summary" in text_lower:
                    current_section = "summary"
                elif "skill" in text_lower:
                    current_section = "skills"
                elif "experience" in text_lower:
                    current_section = "experience"
                elif "education" in text_lower:
                    current_section = "education"
                elif "certif" in text_lower or "award" in text_lower:
                    current_section = "certifications"
                elif "project" in text_lower:
                    current_section = "projects"
                continue

            # Populate sections
            if current_section == "summary":
                if sections["summary"]:
                    sections["summary"] += " " + text
                else:
                    sections["summary"] = text

            elif current_section == "skills":
                skill_text += " " + text

            elif current_section == "experience":
                # Check if this is a bullet point (by style or character)
                is_bullet = (
                    text.startswith(("•", "-", "*", "◦")) or
                    (para.style and "bullet" in para.style.name.lower())
                )

                # Detect if it's a company/title or a bullet
                is_job_title = re.match(r'^[A-Z].*\d{4}', text) or any(
                    kw in text for kw in ["–", "-", "to"]
                )

                if is_job_title and not is_bullet:
                    # Store previous entry
                    if current_entry:
                        sections["experience"].append(current_entry)

                    # Start new entry
                    current_entry = {
                        "company": None,
                        "title": None,
                        "dates": None,
                        "bullets": [],
                    }

                    # Try to parse company/title/dates
                    parts = text.split("|")
                    if len(parts) >= 2:
                        current_entry["company"] = parts[0].strip()
                        current_entry["title"] = parts[1].strip()
                    else:
                        current_entry["company"] = text

                    if len(parts) >= 3:
                        current_entry["dates"] = parts[2].strip()

                elif current_entry and is_bullet:
                    # Remove bullet character if present
                    bullet_text = re.sub(r'^[•\-\*◦]\s*', '', text).strip()
                    current_entry["bullets"].append(bullet_text)
                elif current_entry:
                    # Continuation of previous bullet
                    if current_entry["bullets"]:
                        current_entry["bullets"][-1] += " " + text

            elif current_section == "education":
                sections["education"].append(text)

            elif current_section == "certifications":
                sections["certifications"].append(text)

            elif current_section == "projects":
                sections["projects"].append(text)

        # Add final experience entry
        if current_entry:
            sections["experience"].append(current_entry)

        # Parse skills
        if skill_text:
            skills = [s.strip() for s in re.split(r'[,•\n]', skill_text) if s.strip()]
            sections["skills"] = skills

        self.sections = sections
        logger.info(f"Parsed resume: {len(sections['experience'])} roles, {len(sections['skills'])} skills")
        return sections

    def get_all_bullets(self) -> List[str]:
        """Get all bullet points from experience section."""
        bullets = []
        for entry in self.sections.get("experience", []):
            bullets.extend(entry.get("bullets", []))
        return bullets

    def get_document(self) -> Optional[DocxDocument]:
        """Return the loaded Document object for further processing."""
        return self.doc
