import pytest
import tempfile
import os
from unittest.mock import patch, MagicMock
from docx import Document
from app.resume_tailor import ResumeTailor


@pytest.fixture
def sample_resume_file():
    """Create a sample resume docx for testing."""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()

        # Add contact info
        p = doc.add_paragraph()
        p.add_run("Jane Smith\n")
        p.add_run("jane.smith@example.com")

        # Add summary
        doc.add_paragraph("Professional Summary")
        doc.add_paragraph("Software engineer with experience in backend development.")

        # Add skills section
        doc.add_paragraph("Technical Skills")
        doc.add_paragraph("Python, JavaScript, SQL, Git")

        # Add experience
        doc.add_paragraph("Work Experience")
        doc.add_paragraph("TechCorp | Software Engineer | 2020-2024")
        doc.add_paragraph("• Built REST APIs using Python and FastAPI")
        doc.add_paragraph("• Managed databases with PostgreSQL")

        # Save to temp file
        resume_path = os.path.join(tmpdir, "test_resume.docx")
        doc.save(resume_path)

        yield resume_path


@pytest.fixture
def tailor():
    """Create a ResumeTailor with a mock API key."""
    with patch("app.resume_tailor.Anthropic"):
        return ResumeTailor(api_key="test_key", model="claude-sonnet-4-6")


def test_tailor_initialization(tailor):
    """Test ResumeTailor initialization."""
    assert tailor.model == "claude-sonnet-4-6"
    assert tailor.changes == []


def test_identify_editable_sections(tailor, sample_resume_file):
    """Test identification of editable sections in resume."""
    doc = Document(sample_resume_file)
    editable = tailor._identify_editable_sections(doc)

    # Should find some editable sections (bullets, skills, summary)
    assert len(editable) > 0

    # Check that we have different section types
    section_types = set(e["section_type"] for e in editable)
    assert len(section_types) > 0


def test_is_heading(tailor):
    """Test heading detection."""
    doc = Document()

    # Heading
    heading_para = doc.add_heading("Experience", level=1)
    assert tailor._is_heading(heading_para) == True

    # Regular paragraph
    regular_para = doc.add_paragraph("Some regular text")
    assert tailor._is_heading(regular_para) == False


def test_is_date_range(tailor):
    """Test date range detection."""
    # Should detect dates
    assert tailor._is_date_range("2020-2024") == True
    assert tailor._is_date_range("Jan 2020 – Dec 2024") == True
    assert tailor._is_date_range("January 2020 to December 2024") == True

    # Should not detect regular text
    assert tailor._is_date_range("Software engineer") == False


def test_classify_section(tailor):
    """Test section classification."""
    doc = Document()

    # Bullet
    bullet_para = doc.add_paragraph("• Led a team of engineers")
    assert tailor._classify_section(bullet_para) == "bullet"

    # Skill
    skill_para = doc.add_paragraph("Python, JavaScript, SQL")
    assert tailor._classify_section(skill_para) in ["skill", "other"]

    # Regular text
    text_para = doc.add_paragraph("Some description")
    assert tailor._classify_section(text_para) in ["other", "summary"]


def test_generate_diff(tailor):
    """Test diff generation."""
    tailor.changes = [
        {
            "original": "Worked with Python",
            "new": "Built scalable applications using Python",
            "section_type": "bullet",
        },
        {
            "original": "Python, JavaScript",
            "new": "Python, JavaScript, Kubernetes",
            "section_type": "skill",
        },
    ]

    diff = tailor.generate_diff()

    assert len(diff) == 2
    assert diff[0]["type"] == "change"
    assert diff[0]["section"] == "bullet"
    assert "Worked with Python" in diff[0]["before"]


def test_build_rewrite_prompt(tailor):
    """Test prompt building for Claude."""
    original = "Built REST APIs"
    jd_keywords = ["Python", "FastAPI", "API", "backend"]
    match_report = {
        "matched_keywords": ["Python"],
        "missing_keywords": ["FastAPI", "backend"],
    }

    prompt = tailor._build_rewrite_prompt(
        original,
        "Sample JD text",
        jd_keywords,
        match_report,
        "bullet",
    )

    assert "Original text:" in prompt
    assert original in prompt
    assert "Python" in prompt
    assert "FastAPI" in prompt
    assert "fabricate" in prompt or "exaggerate" in prompt


@patch("app.resume_tailor.Anthropic")
def test_rewrite_text_with_mock_api(mock_anthropic_class, tailor):
    """Test text rewriting with mocked API."""
    # Mock the Anthropic client
    mock_client = MagicMock()
    mock_anthropic_class.return_value = mock_client

    # Mock the API response
    mock_response = MagicMock()
    mock_response.content[0].text = "Designed and implemented REST APIs using Python and FastAPI"
    mock_client.messages.create.return_value = mock_response

    # Create a tailor with the mocked client
    tailor = ResumeTailor(api_key="test_key")
    tailor.client = mock_client

    result = tailor._rewrite_text(
        "Built REST APIs",
        "Sample JD",
        ["Python", "FastAPI"],
        {"matched_keywords": ["Python"], "missing_keywords": ["FastAPI"]},
        "bullet",
    )

    assert result == "Designed and implemented REST APIs using Python and FastAPI"
    mock_client.messages.create.assert_called_once()


def test_save_diff_report(tailor):
    """Test saving diff report."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tailor.changes = [
            {
                "original": "Worked with Python",
                "new": "Built scalable applications using Python",
                "section_type": "bullet",
            },
        ]

        report_path = os.path.join(tmpdir, "diff_report.md")
        tailor.save_diff_report(report_path)

        assert os.path.exists(report_path)

        with open(report_path, "r") as f:
            content = f.read()
            assert "Resume Tailoring Diff Report" in content
            assert "Worked with Python" in content
            assert "BULLET" in content


def test_no_changes_by_default(tailor):
    """Test that a new tailor has no changes."""
    assert tailor.changes == []
    diff = tailor.generate_diff()
    assert diff == []
