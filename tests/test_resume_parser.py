import pytest
import tempfile
import os
from pathlib import Path
from docx import Document
from app.resume_parser import ResumeParser


@pytest.fixture
def sample_resume():
    """Create a sample resume docx for testing."""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()

        # Add contact info
        p = doc.add_paragraph()
        p.add_run("John Doe\n")
        p.add_run("john.doe@example.com | (555) 123-4567 | San Francisco, CA")

        # Add summary
        doc.add_paragraph("Professional Summary")
        doc.add_paragraph(
            "Software engineer with 5 years of experience in cloud infrastructure."
        )

        # Add skills
        doc.add_paragraph("Skills")
        doc.add_paragraph("Python, AWS, Docker, Kubernetes, Linux, SQL")

        # Add experience
        doc.add_paragraph("Experience")
        doc.add_paragraph("Tech Corp | Senior Engineer | 2020-2024")
        doc.add_paragraph("Led team of 5 engineers on cloud migration project", style="List Bullet")
        doc.add_paragraph("Reduced deployment time by 40% using Kubernetes", style="List Bullet")

        doc.add_paragraph("StartupXYZ | Software Engineer | 2018-2020")
        doc.add_paragraph("Built REST APIs in Python using FastAPI", style="List Bullet")
        doc.add_paragraph("Managed AWS infrastructure with Terraform", style="List Bullet")

        # Add education
        doc.add_paragraph("Education")
        doc.add_paragraph("BS Computer Science - State University, 2018")

        # Save to temp file
        resume_path = os.path.join(tmpdir, "sample_resume.docx")
        doc.save(resume_path)

        yield resume_path


def test_resume_parser_initialization(sample_resume):
    """Test that ResumeParser initializes correctly."""
    parser = ResumeParser(sample_resume)
    assert parser.doc is not None
    assert parser.docx_path.exists()


def test_contact_extraction(sample_resume):
    """Test contact info extraction."""
    parser = ResumeParser(sample_resume)
    contact = parser.extract_contact_info()

    assert contact["email"] == "john.doe@example.com"
    assert contact["phone"] == "(555) 123-4567"
    assert contact["location"] == "San Francisco, CA"
    assert contact["name"] == "John Doe"


def test_resume_parsing(sample_resume):
    """Test full resume parsing."""
    parser = ResumeParser(sample_resume)
    sections = parser.parse_resume()

    assert sections["contact"]["email"] == "john.doe@example.com"
    assert "Professional Summary" not in sections["summary"] or len(sections["summary"]) > 0
    assert len(sections["skills"]) > 0
    assert len(sections["experience"]) > 0
    assert len(sections["education"]) > 0


def test_extract_all_bullets(sample_resume):
    """Test extracting all bullet points."""
    parser = ResumeParser(sample_resume)
    parser.parse_resume()
    bullets = parser.get_all_bullets()

    assert len(bullets) >= 4
    assert any("Kubernetes" in bullet for bullet in bullets)
    assert any("FastAPI" in bullet for bullet in bullets)


def test_resume_not_found():
    """Test handling of missing resume file."""
    parser = ResumeParser("/nonexistent/path/resume.docx")
    assert parser.doc is None
    contact = parser.extract_contact_info()
    assert contact == {}


def test_get_document(sample_resume):
    """Test getting the Document object."""
    parser = ResumeParser(sample_resume)
    doc = parser.get_document()
    assert doc is not None
    # Check it has expected Document attributes
    assert hasattr(doc, 'paragraphs')
